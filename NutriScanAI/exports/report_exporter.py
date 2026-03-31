# exports/report_exporter.py
# Generates a downloadable .docx report from report_data + plain text

import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_docx(report_data: dict, plain_text_report: str) -> bytes:
    """
    Generate a formatted .docx report from the report_data dict.
    Returns raw bytes for st.download_button.
    """
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    # ── Title ─────────────────────────────────────────────────────────────────
    title = doc.add_heading("NutriScan AI — Health Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_color(title.runs[0], 0x1B, 0x4F, 0x72)

    subtitle = doc.add_paragraph(
        f"Generated: {report_data.get('generated_at', '')}  |  "
        f"NutriScan AI v2.0"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(10)
    subtitle.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()

    # ── Product info ──────────────────────────────────────────────────────────
    doc.add_heading("Product Information", level=1)

    info_fields = [
        ("Product",     report_data.get("product_name", "Unknown")),
        ("Barcode",     report_data.get("barcode")),
        ("Nutri-Score", report_data.get("nutri_score")),
        ("NOVA Group",  report_data.get("nova_group")),
    ]
    for label, value in info_fields:
        if value:
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(str(value))

    # ── Scores ────────────────────────────────────────────────────────────────
    doc.add_heading("Health Score", level=1)

    score_p = doc.add_paragraph()
    score_p.add_run("Overall Health Score: ").bold = True
    score_p.add_run(f"{report_data.get('display_score', 'N/A')}/10")

    freq_p = doc.add_paragraph()
    freq_p.add_run("Recommended Frequency: ").bold = True
    freq_p.add_run(str(report_data.get("consumption_frequency", "N/A")).title())

    if report_data.get("is_personalized"):
        pers_p = doc.add_paragraph()
        pers_p.add_run("Personalised: ").bold = True
        pers_p.add_run(
            f"General {report_data.get('general_score')}/10 → "
            f"Adjusted {report_data.get('personalized_score')}/10"
        )

    # ── Nutrition table ───────────────────────────────────────────────────────
    nutrition_table = report_data.get("nutrition_table", [])
    if nutrition_table:
        doc.add_heading("Nutritional Values (per 100g)", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"

        hdr = table.rows[0].cells
        for i, text in enumerate(["Nutrient", "Value", "Unit"]):
            hdr[i].text = text
            hdr[i].paragraphs[0].runs[0].bold = True

        for row in nutrition_table:
            cells = table.add_row().cells
            cells[0].text = row["label"]
            cells[1].text = str(row["value"])
            cells[2].text = row["unit"]

    # ── Red flags ─────────────────────────────────────────────────────────────
    red_flags = report_data.get("red_flags", [])
    if red_flags:
        doc.add_heading("Red Flags", level=1)
        for flag in red_flags:
            p = doc.add_paragraph(f"• {flag}", style="List Bullet")

    # ── Positives ─────────────────────────────────────────────────────────────
    positives = report_data.get("positives", [])
    if positives:
        doc.add_heading("Positives", level=1)
        for pos in positives:
            doc.add_paragraph(f"• {pos}", style="List Bullet")

    # ── Personalised section ──────────────────────────────────────────────────
    if report_data.get("is_personalized"):
        doc.add_heading("Personalised Insights", level=1)

        if report_data.get("personalization_note"):
            doc.add_paragraph(report_data["personalization_note"])

        warnings = report_data.get("personalized_warnings", [])
        if warnings:
            doc.add_heading("Warnings for Your Profile", level=2)
            for w in warnings:
                doc.add_paragraph(f"• {w}", style="List Bullet")

        tips = report_data.get("personalized_tips", [])
        if tips:
            doc.add_heading("Tips for You", level=2)
            for t in tips:
                doc.add_paragraph(f"• {t}", style="List Bullet")

    # ── Allergens & additives ─────────────────────────────────────────────────
    allergens = report_data.get("allergens", [])
    additives = report_data.get("additives_flagged", [])
    if allergens or additives:
        doc.add_heading("Allergens & Additives", level=1)
        if allergens:
            p = doc.add_paragraph()
            p.add_run("Allergens: ").bold = True
            p.add_run(", ".join(allergens))
        if additives:
            p = doc.add_paragraph()
            p.add_run("Flagged Additives: ").bold = True
            p.add_run(", ".join(additives))

    # ── Narrative ─────────────────────────────────────────────────────────────
    narrative = report_data.get("narrative") or report_data.get("health_summary", "")
    if narrative:
        doc.add_heading("Health Report Summary", level=1)
        doc.add_paragraph(narrative)

    # ── Disclaimer ────────────────────────────────────────────────────────────
    doc.add_paragraph()
    disclaimer = doc.add_paragraph(
        "This report is generated by NutriScan AI for informational purposes only. "
        "It does not constitute medical advice. Please consult a qualified dietitian "
        "or healthcare professional for personalised dietary guidance."
    )
    disclaimer.runs[0].font.size = Pt(9)
    disclaimer.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    disclaimer.runs[0].font.italic = True

    # ── Serialise to bytes ────────────────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _set_run_color(run, r: int, g: int, b: int):
    run.font.color.rgb = RGBColor(r, g, b)